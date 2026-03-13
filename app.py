#!/usr/bin/env python3
"""
格不丢翻译 API - Flask 后端
真正的格式保留翻译服务
"""

import os
import hashlib
import json
import time
from io import BytesIO
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app, origins="*")

# 7大专业领域
DOMAINS = {
    "general": {
        "name": "通用领域",
        "prompt": "Translate this Chinese text to English for general business use. Maintain professional tone."
    },
    "electronics": {
        "name": "电子产品",
        "prompt": "Translate this Chinese electronics specification to English. Preserve model numbers, technical terms, and units exactly. Use industry-standard terminology."
    },
    "medical": {
        "name": "医疗器械",
        "prompt": "Translate this Chinese medical device content to English. Use formal, cautious medical terminology. Ensure accuracy for regulatory compliance."
    },
    "legal": {
        "name": "法律合同",
        "prompt": "Translate this Chinese legal document to English. Use precise legal terminology. Maintain formal contractual language and obligation expressions."
    },
    "marketing": {
        "name": "市场营销",
        "prompt": "Translate this Chinese marketing copy to English. Make it persuasive and appealing. Adapt for cultural resonance while maintaining brand voice."
    },
    "industrial": {
        "name": "工业技术",
        "prompt": "Translate this Chinese industrial/technical document to English. Use ISO-standard terminology. Be precise with technical specifications."
    },
    "software": {
        "name": "软件/IT",
        "prompt": "Translate this Chinese software/IT documentation to English. Preserve code snippets, API names, and technical terms. Use developer-friendly language."
    }
}

# DeepSeek 客户端
deepseek_client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY", "sk-37a56b14534e450fbe6068c95cff4044"),
    base_url="https://api.deepseek.com/v1"
)

# 缓存
translation_cache = {}
CACHE_FILE = ".translation_cache.json"

# 加载缓存
try:
    with open(CACHE_FILE, 'r', encoding='utf-8') as f:
        translation_cache = json.load(f)
except:
    pass

def get_cache_key(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def should_translate(text):
    """判断文本是否需要翻译"""
    if not text or len(text.strip()) < 2:
        return False
    text = text.strip()
    # 跳过纯数字、URL、邮箱、型号等
    skip_patterns = [
        lambda x: x.isdigit(),
        lambda x: x.startswith('http'),
        lambda x: '@' in x and '.' in x.split('@')[-1],
        lambda x: x.startswith('G-') and x[2:].replace('-', '').isalnum(),
        lambda x: x.replace('.', '').replace('-', '').isdigit(),
    ]
    return not any(p(text) for p in skip_patterns)

def translate_text(text, domain="general"):
    """翻译单个文本段"""
    if not should_translate(text):
        return text
    
    # 检查缓存
    cache_key = get_cache_key(f"{domain}:{text}")
    if cache_key in translation_cache:
        return translation_cache[cache_key]
    
    # 调用 DeepSeek API
    domain_config = DOMAINS.get(domain, DOMAINS["general"])
    
    try:
        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {
                    "role": "system",
                    "content": f"You are a professional translator. {domain_config['prompt']} Provide ONLY the English translation, no explanations."
                },
                {
                    "role": "user",
                    "content": f"Translate to English:\n{text}"
                }
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        translated = response.choices[0].message.content.strip()
        
        # 保存缓存
        translation_cache[cache_key] = translated
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(translation_cache, f, ensure_ascii=False, indent=2)
        
        return translated
    except Exception as e:
        print(f"Translation error: {e}")
        return text

def translate_docx(file_buffer, domain="general"):
    """翻译 DOCX 文件，保留格式"""
    doc = Document(BytesIO(file_buffer))
    
    # 统计信息
    stats = {"paragraphs": 0, "tables": 0, "cells": 0}
    
    # 翻译段落
    for para in doc.paragraphs:
        if para.text.strip():
            stats["paragraphs"] += 1
            # 合并所有 runs 的文本
            full_text = para.text
            translated = translate_text(full_text, domain)
            
            # 保留第一个 run 的格式，清空其他 runs
            if para.runs:
                para.runs[0].text = translated
                for run in para.runs[1:]:
                    run.text = ""
    
    # 翻译表格
    for table in doc.tables:
        stats["tables"] += 1
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    stats["cells"] += 1
                    # 翻译单元格内的段落
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text = para.text
                            translated = translate_text(full_text, domain)
                            if para.runs:
                                para.runs[0].text = translated
                                for run in para.runs[1:]:
                                    run.text = ""
    
    # 翻译页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        translated = translate_text(para.text, domain)
                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""
        
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    if para.text.strip():
                        translated = translate_text(para.text, domain)
                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, stats

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        "status": "ok",
        "service": "格不丢翻译 API",
        "version": "1.0.0",
        "domains": list(DOMAINS.keys())
    })

@app.route('/domains', methods=['GET'])
def get_domains():
    """获取支持的领域列表"""
    return jsonify({
        "domains": {k: v["name"] for k, v in DOMAINS.items()}
    })

@app.route('/translate', methods=['POST'])
def translate():
    """翻译文件"""
    start_time = time.time()
    
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    domain = request.form.get('domain', 'general')
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    # 检查文件格式
    is_docx = file.filename.endswith('.docx')
    is_xlsx = file.filename.endswith('.xlsx')
    
    if not is_docx and not is_xlsx:
        return jsonify({"error": "Only .docx and .xlsx files are supported"}), 400
    
    # 读取文件
    file_buffer = file.read()
    file_size = len(file_buffer)
    
    # 检查文件大小 (16MB limit)
    if file_size > 16 * 1024 * 1024:
        return jsonify({"error": "File too large. Maximum size is 16MB"}), 413
    
    try:
        if is_docx:
            output, stats = translate_docx(file_buffer, domain)
            output_filename = file.filename.replace('.docx', '_EN.docx')
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            # Excel 翻译 (简化版)
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(file_buffer))
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            if should_translate(cell.value):
                                cell.value = translate_text(cell.value, domain)
            
            output = BytesIO()
            wb.save(output)
            output.seek(0)
            stats = {"sheets": len(wb.sheetnames)}
            output_filename = file.filename.replace('.xlsx', '_EN.xlsx')
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        processing_time = time.time() - start_time
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=output_filename
        )
    
    except Exception as e:
        print(f"Translation error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
