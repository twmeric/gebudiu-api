#!/usr/bin/env python3
"""
========================================
格不丢翻译 API - 重构版
GeBuDiu Translation API - Refactored
========================================

主要改进：
1. 模块化代码结构
2. 完善的错误处理
3. 日志记录
4. 配置管理
5. 缓存优化

技术栈：
- Flask: Web 框架
- OpenAI: DeepSeek API 调用
- python-docx: Word 文档处理
- openpyxl: Excel 文档处理
"""

import os
import sys
import hashlib
import json
import time
import logging
from io import BytesIO
from datetime import datetime
from functools import wraps
from typing import Dict, List, Tuple, Optional, Union

from flask import Flask, request, send_file, jsonify, Response
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI, APIError as OpenAIError
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# ============================================
# 配置管理
# ============================================

class Config:
    """应用配置类"""
    
    # API 配置
    DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
    DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"
    
    # 服务配置
    DEBUG = os.getenv("DEBUG", "False").lower() == "true"
    PORT = int(os.getenv("PORT", 5000))
    HOST = os.getenv("HOST", "0.0.0.0")
    
    # 文件限制
    MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
    ALLOWED_EXTENSIONS = {'.docx', '.xlsx'}
    
    # 缓存配置
    CACHE_FILE = ".translation_cache.json"
    CACHE_MAX_SIZE = 10000  # 最大缓存条目数
    
    # 日志配置
    LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")
    LOG_FORMAT = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'


# ============================================
# 日志设置
# ============================================

def setup_logging():
    """配置日志记录"""
    logging.basicConfig(
        level=getattr(logging, Config.LOG_LEVEL),
        format=Config.LOG_FORMAT,
        handlers=[
            logging.StreamHandler(sys.stdout),
            logging.FileHandler('gebudiu.log', encoding='utf-8')
        ]
    )
    return logging.getLogger('gebudiu')

logger = setup_logging()


# ============================================
# 领域配置
# ============================================

class DomainConfig:
    """专业领域配置"""
    
    DOMAINS: Dict[str, Dict[str, str]] = {
        "general": {
            "name": "通用领域",
            "prompt": "Translate this Chinese text to English for general business use. Maintain professional tone."
        },
        "electronics": {
            "name": "电子产品",
            "prompt": "Translate this Chinese electronics specification to English. Preserve model numbers, technical terms, and units exactly. Use industry-standard terminology."
        },
        "medical": {
            "name": "医疗器械",
            "prompt": "Translate this Chinese medical device content to English. Use formal, cautious medical terminology. Ensure accuracy for regulatory compliance."
        },
        "legal": {
            "name": "法律合同",
            "prompt": "Translate this Chinese legal document to English. Use precise legal terminology. Maintain formal contractual language and obligation expressions."
        },
        "marketing": {
            "name": "市场营销",
            "prompt": "Translate this Chinese marketing copy to English. Make it persuasive and appealing. Adapt for cultural resonance while maintaining brand voice."
        },
        "industrial": {
            "name": "工业技术",
            "prompt": "Translate this Chinese industrial/technical document to English. Use ISO-standard terminology. Be precise with technical specifications."
        },
        "software": {
            "name": "软件/IT",
            "prompt": "Translate this Chinese software/IT documentation to English. Preserve code snippets, API names, and technical terms. Use developer-friendly language."
        }
    }
    
    @classmethod
    def get(cls, domain: str) -> Dict[str, str]:
        """获取领域配置"""
        return cls.DOMAINS.get(domain, cls.DOMAINS["general"])
    
    @classmethod
    def get_name(cls, domain: str) -> str:
        """获取领域名称"""
        return cls.get(domain)["name"]
    
    @classmethod
    def get_prompt(cls, domain: str) -> str:
        """获取领域提示词"""
        return cls.get(domain)["prompt"]
    
    @classmethod
    def list_domains(cls) -> Dict[str, str]:
        """获取所有领域列表"""
        return {k: v["name"] for k, v in cls.DOMAINS.items()}


# ============================================
# 缓存管理
# ============================================

class TranslationCache:
    """翻译缓存管理器"""
    
    def __init__(self, cache_file: str = Config.CACHE_FILE):
        self.cache_file = cache_file
        self.cache: Dict[str, str] = {}
        self._load()
    
    def _load(self):
        """从文件加载缓存"""
        try:
            if os.path.exists(self.cache_file):
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self.cache = json.load(f)
                logger.info(f"Loaded {len(self.cache)} cached translations")
        except Exception as e:
            logger.warning(f"Failed to load cache: {e}")
            self.cache = {}
    
    def _save(self):
        """保存缓存到文件"""
        try:
            # 限制缓存大小
            if len(self.cache) > Config.CACHE_MAX_SIZE:
                # 保留最新的条目
                items = list(self.cache.items())
                self.cache = dict(items[-Config.CACHE_MAX_SIZE:])
            
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"Failed to save cache: {e}")
    
    def _make_key(self, text: str, domain: str) -> str:
        """生成缓存键"""
        return hashlib.md5(f"{domain}:{text}".encode('utf-8')).hexdigest()
    
    def get(self, text: str, domain: str) -> Optional[str]:
        """获取缓存的翻译"""
        key = self._make_key(text, domain)
        return self.cache.get(key)
    
    def set(self, text: str, domain: str, translation: str):
        """设置缓存"""
        key = self._make_key(text, domain)
        self.cache[key] = translation
        self._save()
    
    def clear(self):
        """清空缓存"""
        self.cache.clear()
        self._save()


# 全局缓存实例
cache = TranslationCache()


# ============================================
# 翻译服务
# ============================================

class TranslationService:
    """翻译服务核心类"""
    
    def __init__(self):
        self.client = OpenAI(
            api_key=Config.DEEPSEEK_API_KEY,
            base_url=Config.DEEPSEEK_BASE_URL
        )
        self.model = "deepseek-chat"
        self.temperature = 0.3
        self.max_tokens = 1000
    
    def _should_translate(self, text: str) -> bool:
        """
        判断文本是否需要翻译
        
        规则：
        - 空文本或太短的不翻译
        - 纯数字、URL、邮箱、型号不翻译
        """
        if not text or len(text.strip()) < 2:
            return False
        
        text = text.strip()
        
        # 跳过模式
        skip_patterns = [
            lambda x: x.isdigit(),  # 纯数字
            lambda x: x.startswith('http'),  # URL
            lambda x: '@' in x and '.' in x.split('@')[-1],  # 邮箱
            lambda x: x.startswith('G-') and x[2:].replace('-', '').isalnum(),  # 型号
            lambda x: x.replace('.', '').replace('-', '').isdigit(),  # 数字组合
        ]
        
        return not any(p(text) for p in skip_patterns)
    
    def translate_text(self, text: str, domain: str = "general") -> str:
        """
        翻译单个文本段
        
        Args:
            text: 要翻译的文本
            domain: 专业领域
            
        Returns:
            翻译后的文本
        """
        # 检查是否需要翻译
        if not self._should_translate(text):
            return text
        
        # 检查缓存
        cached = cache.get(text, domain)
        if cached:
            logger.debug(f"Cache hit for text: {text[:50]}...")
            return cached
        
        # 调用 API 翻译
        try:
            prompt = DomainConfig.get_prompt(domain)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. {prompt} Provide ONLY the English translation, no explanations."
                    },
                    {
                        "role": "user",
                        "content": f"Translate to English:\n{text}"
                    }
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens
            )
            
            translated = response.choices[0].message.content.strip()
            
            # 保存缓存
            cache.set(text, domain, translated)
            
            logger.debug(f"Translated: {text[:50]}... -> {translated[:50]}...")
            return translated
            
        except OpenAIError as e:
            logger.error(f"OpenAI API error: {e}")
            raise TranslationError(f"翻译服务暂时不可用: {str(e)}")
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return text  # 出错时返回原文


class TranslationError(Exception):
    """翻译错误异常"""
    pass


# 全局翻译服务实例
translation_service = TranslationService()


# ============================================
# 文档处理器
# ============================================

class DocumentProcessor:
    """文档处理基类"""
    
    def __init__(self, service: TranslationService):
        self.service = service
    
    def process(self, file_buffer: bytes, domain: str) -> Tuple[BytesIO, Dict]:
        """处理文档，子类必须实现"""
        raise NotImplementedError


class DocxProcessor(DocumentProcessor):
    """Word 文档处理器"""
    
    def process(self, file_buffer: bytes, domain: str) -> Tuple[BytesIO, Dict]:
        """
        翻译 DOCX 文件，保留格式
        
        Returns:
            (输出缓冲区, 统计信息)
        """
        doc = Document(BytesIO(file_buffer))
        stats = {"paragraphs": 0, "tables": 0, "cells": 0, "headers": 0, "footers": 0}
        
        # 翻译段落
        for para in doc.paragraphs:
            if para.text.strip():
                stats["paragraphs"] += 1
                translated = self.service.translate_text(para.text, domain)
                
                # 保留第一个 run 的格式，清空其他 runs
                if para.runs:
                    para.runs[0].text = translated
                    for run in para.runs[1:]:
                        run.text = ""
        
        # 翻译表格
        for table in doc.tables:
            stats["tables"] += 1
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        stats["cells"] += 1
                        for para in cell.paragraphs:
                            if para.text.strip():
                                translated = self.service.translate_text(para.text, domain)
                                if para.runs:
                                    para.runs[0].text = translated
                                    for run in para.runs[1:]:
                                        run.text = ""
        
        # 翻译页眉页脚
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            stats["headers"] += 1
                            translated = self.service.translate_text(para.text, domain)
                            if para.runs:
                                para.runs[0].text = translated
                                for run in para.runs[1:]:
                                    run.text = ""
            
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            stats["footers"] += 1
                            translated = self.service.translate_text(para.text, domain)
                            if para.runs:
                                para.runs[0].text = translated
                                for run in para.runs[1:]:
                                    run.text = ""
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output, stats


class XlsxProcessor(DocumentProcessor):
    """Excel 文档处理器"""
    
    def process(self, file_buffer: bytes, domain: str) -> Tuple[BytesIO, Dict]:
        """
        翻译 XLSX 文件
        
        Returns:
            (输出缓冲区, 统计信息)
        """
        from openpyxl import load_workbook
        
        wb = load_workbook(BytesIO(file_buffer))
        stats = {"sheets": 0, "cells_translated": 0}
        
        for sheet_name in wb.sheetnames:
            stats["sheets"] += 1
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        if self.service._should_translate(cell.value):
                            cell.value = self.service.translate_text(cell.value, domain)
                            stats["cells_translated"] += 1
        
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        return output, stats


# 处理器映射
PROCESSORS = {
    '.docx': DocxProcessor,
    '.xlsx': XlsxProcessor
}


# ============================================
# Flask 应用
# ============================================

def create_app() -> Flask:
    """创建 Flask 应用实例"""
    app = Flask(__name__)
    CORS(app, origins="*")
    
    # ============================================
    # 错误处理
    # ============================================
    
    @app.errorhandler(400)
    def bad_request(error):
        logger.warning(f"Bad request: {error}")
        return jsonify({"error": "请求参数错误", "details": str(error)}), 400
    
    @app.errorhandler(413)
    def too_large(error):
        logger.warning(f"File too large: {error}")
        return jsonify({"error": f"文件太大，最大支持 {Config.MAX_FILE_SIZE // 1024 // 1024}MB"}), 413
    
    @app.errorhandler(500)
    def server_error(error):
        logger.error(f"Server error: {error}")
        return jsonify({"error": "服务器内部错误", "details": str(error)}), 500
    
    @app.errorhandler(TranslationError)
    def translation_error(error):
        logger.error(f"Translation error: {error}")
        return jsonify({"error": str(error)}), 500
    
    # ============================================
    # 路由
    # ============================================
    
    @app.route('/health', methods=['GET'])
    def health() -> Response:
        """健康检查接口"""
        return jsonify({
            "status": "ok",
            "service": "格不丢翻译 API",
            "version": "2.0.0",
            "domains": list(DomainConfig.DOMAINS.keys()),
            "timestamp": datetime.now().isoformat()
        })
    
    @app.route('/domains', methods=['GET'])
    def get_domains() -> Response:
        """获取支持的领域列表"""
        return jsonify({
            "domains": DomainConfig.list_domains()
        })
    
    @app.route('/translate', methods=['POST'])
    def translate() -> Response:
        """翻译文件接口"""
        start_time = time.time()
        
        # 检查文件
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        domain = request.form.get('domain', 'general')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # 检查文件格式
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in Config.ALLOWED_EXTENSIONS:
            return jsonify({
                "error": f"Unsupported file format. Allowed: {', '.join(Config.ALLOWED_EXTENSIONS)}"
            }), 400
        
        # 读取文件
        file_buffer = file.read()
        file_size = len(file_buffer)
        
        # 检查文件大小
        if file_size > Config.MAX_FILE_SIZE:
            return jsonify({
                "error": f"File too large ({file_size // 1024 // 1024}MB). Maximum size is {Config.MAX_FILE_SIZE // 1024 // 1024}MB"
            }), 413
        
        logger.info(f"Processing file: {file.filename}, size: {file_size}, domain: {domain}")
        
        try:
            # 获取处理器
            processor_class = PROCESSORS.get(file_ext)
            if not processor_class:
                return jsonify({"error": "No processor available for this file type"}), 500
            
            processor = processor_class(translation_service)
            output, stats = processor.process(file_buffer, domain)
            
            # 生成输出文件名
            output_filename = file.filename.rsplit('.', 1)[0] + '_EN' + file_ext
            
            # 设置 MIME 类型
            mimetypes = {
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
            
            processing_time = time.time() - start_time
            logger.info(f"Translation completed in {processing_time:.2f}s. Stats: {stats}")
            
            return send_file(
                output,
                mimetype=mimetypes.get(file_ext, 'application/octet-stream'),
                as_attachment=True,
                download_name=output_filename
            )
        
        except Exception as e:
            logger.exception("Translation failed")
            return jsonify({"error": f"Translation failed: {str(e)}"}), 500
    
    @app.route('/cache/clear', methods=['POST'])
    def clear_cache() -> Response:
        """清空翻译缓存（管理员接口）"""
        # 这里可以添加认证检查
        cache.clear()
        logger.info("Translation cache cleared")
        return jsonify({"message": "Cache cleared successfully"})
    
    return app


# ============================================
# 入口点
# ============================================

app = create_app()

if __name__ == '__main__':
    logger.info(f"Starting GeBuDiu Translation API on {Config.HOST}:{Config.PORT}")
    logger.info(f"Debug mode: {Config.DEBUG}")
    app.run(debug=Config.DEBUG, host=Config.HOST, port=Config.PORT)
