#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 佈局分析器
自動分析原文和譯文的佈局差異，無需用戶反饋
"""

import logging
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional
from io import BytesIO
import docx
from docx.shared import Pt

from format_fingerprint import ContentFingerprint, FormatOutcome, FormatDiffReport, FormatParams

logger = logging.getLogger(__name__)

@dataclass
class LayoutMetrics:
    """詳細的佈局指標"""
    # 基本統計
    total_chars: int
    paragraph_count: int
    sentence_count: int
    word_count: int
    
    # 格式統計
    avg_font_size: float
    line_spacing: float
    paragraph_spacing: float
    
    # 結構統計
    table_count: int
    image_count: int
    heading_count: int
    list_count: int
    
    # 佈局估算
    estimated_pages: int
    content_density: float  # 內容密度 (0-1)
    
    def to_fingerprint(self, domain: str = "general") -> ContentFingerprint:
        """轉換為內容指紋"""
        avg_sent_len = self.total_chars / max(self.sentence_count, 1)
        complexity = self._calculate_complexity()
        
        return ContentFingerprint(
            domain=domain,
            total_chars=self.total_chars,
            avg_sentence_length=avg_sent_len,
            paragraph_count=self.paragraph_count,
            table_count=self.table_count,
            image_count=self.image_count,
            structure_complexity=complexity
        )
    
    def _calculate_complexity(self) -> float:
        """計算結構複雜度 (0-1)"""
        score = 0.0
        
        # 表格增加複雜度
        score += min(self.table_count * 0.05, 0.3)
        
        # 圖片增加複雜度
        score += min(self.image_count * 0.03, 0.2)
        
        # 標題層次
        score += min(self.heading_count * 0.02, 0.2)
        
        # 列表
        score += min(self.list_count * 0.02, 0.15)
        
        return min(score, 1.0)


class DocxLayoutAnalyzer:
    """DOCX 文件佈局分析器"""
    
    # 估算參數 (A4 頁面)
    CHARS_PER_PAGE_FULL = 1500      # 滿版文字
    CHARS_PER_PAGE_NORMAL = 1200    # 正常邊距
    CHARS_PER_PAGE_SPARSE = 900     # 稀疏佈局
    
    def __init__(self):
        self.metrics = None
    
    def analyze(self, doc_bytes: bytes) -> LayoutMetrics:
        """分析 DOCX 文件佈局"""
        try:
            doc = docx.Document(BytesIO(doc_bytes))
            return self._extract_metrics(doc)
        except Exception as e:
            logger.error(f"Failed to analyze DOCX: {e}")
            # 返回默認值
            return LayoutMetrics(
                total_chars=0, paragraph_count=0, sentence_count=0, word_count=0,
                avg_font_size=11.0, line_spacing=1.15, paragraph_spacing=6.0,
                table_count=0, image_count=0, heading_count=0, list_count=0,
                estimated_pages=1, content_density=0.5
            )
    
    def _extract_metrics(self, doc: docx.Document) -> LayoutMetrics:
        """從 Document 對象提取指標"""
        total_chars = 0
        paragraph_count = 0
        sentence_count = 0
        word_count = 0
        font_sizes = []
        heading_count = 0
        list_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            paragraph_count += 1
            total_chars += len(text)
            
            # 句子數估算 (簡易算法)
            sentence_count += text.count('。') + text.count('.') + text.count('!') + text.count('?')
            sentence_count = max(sentence_count, paragraph_count)  # 至少每段一句
            
            # 詞數估算
            word_count += len(text.split())
            
            # 檢測標題
            if para.style.name.startswith('Heading'):
                heading_count += 1
            
            # 檢測列表
            if para.style.name.startswith('List'):
                list_count += 1
            
            # 提取字體大小
            for run in para.runs:
                if run.font.size:
                    font_sizes.append(run.font.size.pt)
        
        # 表格
        table_count = len(doc.tables)
        
        # 圖片
        image_count = len(doc.inline_shapes)
        
        # 計算平均值
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 11.0
        
        # 估算頁數
        estimated_pages = self._estimate_pages(total_chars, paragraph_count, table_count)
        
        # 內容密度
        content_density = min(total_chars / (estimated_pages * self.CHARS_PER_PAGE_FULL), 1.0)
        
        return LayoutMetrics(
            total_chars=total_chars,
            paragraph_count=paragraph_count,
            sentence_count=max(sentence_count, 1),
            word_count=word_count,
            avg_font_size=avg_font_size,
            line_spacing=1.15,  # 默認值
            paragraph_spacing=6.0,  # 默認值
            table_count=table_count,
            image_count=image_count,
            heading_count=heading_count,
            list_count=list_count,
            estimated_pages=estimated_pages,
            content_density=content_density
        )
    
    def _estimate_pages(self, chars: int, paragraphs: int, tables: int) -> int:
        """估算頁數"""
        # 基於字符數
        pages_from_chars = max(1, chars / self.CHARS_PER_PAGE_NORMAL)
        
        # 表格通常佔用更多空間
        table_overhead = tables * 0.5  # 每個表格額外半頁
        
        return int(pages_from_chars + table_overhead + 0.5)
    
    def compare(self, source_metrics: LayoutMetrics, target_metrics: LayoutMetrics) -> FormatDiffReport:
        """
        比較原文和譯文的佈局差異
        生成客觀的改進建議
        """
        # 計算關鍵指標
        text_expansion_ratio = target_metrics.total_chars / max(source_metrics.total_chars, 1)
        paragraph_diff = target_metrics.paragraph_count - source_metrics.paragraph_count
        paragraph_growth_rate = paragraph_diff / max(source_metrics.paragraph_count, 1)
        font_size_diff = target_metrics.avg_font_size - source_metrics.avg_font_size
        page_increase = target_metrics.estimated_pages - source_metrics.estimated_pages
        
        # 估算留白佔比
        whitespace_ratio = 1 - target_metrics.content_density
        
        # 斷行密度 (簡易估算)
        line_break_density = paragraph_growth_rate * 0.3
        
        # 判斷嚴重程度
        severity = "minor"
        if text_expansion_ratio > 1.5 or page_increase > 5:
            severity = "major"
        elif text_expansion_ratio > 1.3 or page_increase > 2:
            severity = "moderate"
        
        # 生成建議
        suggestions = self._generate_suggestions(
            text_expansion_ratio, page_increase, whitespace_ratio, 
            paragraph_growth_rate, source_metrics.domain
        )
        
        # 計算滿意度
        outcome = FormatOutcome(
            page_count=target_metrics.estimated_pages,
            text_expansion_ratio=text_expansion_ratio,
            paragraph_growth_rate=paragraph_growth_rate,
            whitespace_ratio=whitespace_ratio,
            line_break_density=line_break_density
        )
        satisfaction_score = outcome.calculate_satisfaction_score()
        
        # 自動修復參數
        auto_fix_params = self._calculate_optimal_params(
            text_expansion_ratio, target_metrics
        )
        
        return FormatDiffReport(
            text_expansion_ratio=text_expansion_ratio,
            paragraph_diff=paragraph_diff,
            font_size_diff=font_size_diff,
            page_increase=page_increase,
            severity=severity,
            satisfaction_score=satisfaction_score,
            suggestions=suggestions,
            auto_fix_params=auto_fix_params
        )
    
    def _generate_suggestions(self, expansion_ratio: float, page_increase: int,
                             whitespace_ratio: float, para_growth: float,
                             domain: str) -> List[str]:
        """基於差異生成改進建議"""
        suggestions = []
        
        if expansion_ratio > 1.4:
            suggestions.append(f"🚨 文本膨脹嚴重 ({expansion_ratio:.1%})，建議字體縮小至 10pt")
        elif expansion_ratio > 1.25:
            suggestions.append(f"⚠️ 文本膨脹中等 ({expansion_ratio:.1%})，建議字體調整為 10.5pt")
        
        if page_increase > 3:
            suggestions.append(f"📄 頁數增加 {page_increase} 頁，建議緊湊佈局")
        
        if whitespace_ratio > 0.35:
            suggestions.append(f"⬜ 留白佔比較高 ({whitespace_ratio:.1%})，建議縮小頁邊距")
        
        if para_growth > 0.2:
            suggestions.append(f"📝 段落增加過多 ({para_growth:.1%})，建議檢查斷行邏輯")
        
        if domain in ['medical', 'legal'] and expansion_ratio > 1.2:
            suggestions.append("⚖️ 專業領域文件建議保持標準字體，通過調整行距控制頁數")
        
        if not suggestions:
            suggestions.append("✅ 佈局表現良好，當前參數適合此類內容")
        
        return suggestions
    
    def _calculate_optimal_params(self, expansion_ratio: float, 
                                  metrics: LayoutMetrics) -> FormatParams:
        """計算最佳格式參數"""
        params = FormatParams()
        
        # 基於文本膨脹率調整字體
        if expansion_ratio > 1.5:
            params.font_size = 9.5
            params.line_spacing = 1.0
            params.paragraph_spacing = 3.0
            params.margin_cm = 2.0
        elif expansion_ratio > 1.3:
            params.font_size = 10.0
            params.line_spacing = 1.05
            params.paragraph_spacing = 4.0
            params.margin_cm = 2.2
        elif expansion_ratio > 1.15:
            params.font_size = 10.5
            params.line_spacing = 1.1
        else:
            params.font_size = 11.0
            params.line_spacing = 1.15
        
        # 多表格時進一步緊湊
        if metrics.table_count > 3:
            params.font_size = max(9.0, params.font_size - 0.5)
            params.line_spacing = max(1.0, params.line_spacing - 0.05)
        
        return params


class FormatLearningPipeline:
    """格式學習流水線 - 整合分析和學習"""
    
    def __init__(self, learning_engine):
        self.analyzer = DocxLayoutAnalyzer()
        self.engine = learning_engine
    
    def process_translation(self,
                           source_bytes: bytes,
                           translated_bytes: bytes,
                           domain: str,
                           params_used: FormatParams) -> Dict:
        """
        處理一次翻譯的完整流程：
        1. 分析原文佈局
        2. 分析譯文佈局
        3. 比較差異
        4. 記錄到學習引擎
        """
        # 分析
        source_metrics = self.analyzer.analyze(source_bytes)
        target_metrics = self.analyzer.analyze(translated_bytes)
        
        # 添加領域信息
        source_metrics.domain = domain
        target_metrics.domain = domain
        
        # 比較
        diff_report = self.analyzer.compare(source_metrics, target_metrics)
        
        # 生成指紋並記錄
        fingerprint = source_metrics.to_fingerprint(domain)
        
        outcome = FormatOutcome(
            page_count=target_metrics.estimated_pages,
            text_expansion_ratio=diff_report.text_expansion_ratio,
            paragraph_growth_rate=diff_report.paragraph_diff / max(source_metrics.paragraph_count, 1),
            whitespace_ratio=1 - target_metrics.content_density,
            line_break_density=diff_report.paragraph_diff * 0.01
        )
        
        # 記錄到學習引擎（核心學習步驟）
        self.engine.record_translation_outcome(fingerprint, params_used, outcome)
        
        return {
            "analysis": {
                "source": {
                    "chars": source_metrics.total_chars,
                    "paragraphs": source_metrics.paragraph_count,
                    "estimated_pages": source_metrics.estimated_pages
                },
                "target": {
                    "chars": target_metrics.total_chars,
                    "paragraphs": target_metrics.paragraph_count,
                    "estimated_pages": target_metrics.estimated_pages
                }
            },
            "diff": {
                "text_expansion_ratio": round(diff_report.text_expansion_ratio, 2),
                "page_increase": diff_report.page_increase,
                "severity": diff_report.severity,
                "satisfaction_score": round(diff_report.satisfaction_score, 2)
            },
            "suggestions": diff_report.suggestions,
            "auto_fix_params": diff_report.auto_fix_params.to_dict()
        }
