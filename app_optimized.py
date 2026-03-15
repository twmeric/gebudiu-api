#!/usr/bin/env python3
"""
格不丢翻译 API - Optimized Version
Memory-efficient, batched, cached translation service
"""

import os
import hashlib
import json
import time
import re
import asyncio
from io import BytesIO
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

try:
    from cachetools import TTLCache
    HAS_CACHETOOLS = True
except ImportError:
    HAS_CACHETOOLS = False

from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app, origins="*")

# ============================================================================
# CONFIGURATION
# ============================================================================

BATCH_SIZE = int(os.getenv("BATCH_SIZE", "10"))
MAX_CONCURRENT = int(os.getenv("MAX_CONCURRENT", "5"))
CACHE_SIZE = int(os.getenv("CACHE_SIZE", "10000"))
CACHE_TTL = int(os.getenv("CACHE_TTL", "3600"))  # 1 hour
USE_COMPACT_PROMPTS = os.getenv("USE_COMPACT_PROMPTS", "true").lower() == "true"

# 7大专业领域 - Optimized prompts for token efficiency
DOMAINS = {
    "general": {
        "name": "通用领域",
        "prompt": "Translate Chinese to English. Provide ONLY the translation.",
        "compact": "zh→en:"
    },
    "electronics": {
        "name": "电子产品",
        "prompt": "Translate Chinese electronics spec to English. Preserve model numbers, technical terms. Use industry-standard terminology.",
        "compact": "Electronics zh→en:"
    },
    "medical": {
        "name": "医疗器械",
        "prompt": "Translate Chinese medical content to English. Use formal medical terminology. Ensure regulatory accuracy.",
        "compact": "Medical zh→en:"
    },
    "legal": {
        "name": "法律合同",
        "prompt": "Translate Chinese legal doc to English. Use precise legal terminology. Maintain formal contractual language.",
        "compact": "Legal zh→en:"
    },
    "marketing": {
        "name": "市场营销",
        "prompt": "Translate Chinese marketing to English. Make it persuasive. Adapt for cultural resonance.",
        "compact": "Marketing zh→en:"
    },
    "industrial": {
        "name": "工业技术",
        "prompt": "Translate Chinese industrial doc to English. Use ISO-standard terminology. Be precise with specs.",
        "compact": "Industrial zh→en:"
    },
    "software": {
        "name": "软件/IT",
        "prompt": "Translate Chinese IT doc to English. Preserve code snippets, API names. Use developer-friendly language.",
        "compact": "IT zh→en:"
    }
}

# DeepSeek client with connection pooling
_deepseek_client = None

def get_deepseek_client():
    global _deepseek_client
    if _deepseek_client is None:
        _deepseek_client = OpenAI(
            api_key=os.getenv("DEEPSEEK_API_KEY", ""),
            base_url="https://api.deepseek.com/v1",
            timeout=60.0,
            max_retries=2
        )
    return _deepseek_client

# ============================================================================
# OPTIMIZED CACHING SYSTEM
# ============================================================================

class TranslationCache:
    """Hybrid in-memory + disk cache for translations"""
    
    def __init__(self, cache_dir: str = ".cache", memory_size: int = 10000, ttl: int = 3600):
        self.cache_dir = cache_dir
        os.makedirs(cache_dir, exist_ok=True)
        
        # In-memory LRU cache
        if HAS_CACHETOOLS:
            self.memory_cache = TTLCache(maxsize=memory_size, ttl=ttl)
        else:
            self.memory_cache = {}
            self._memory_keys = []
            self._max_memory = memory_size
        
        self.disk_cache_path = os.path.join(cache_dir, "translations.json")
        self.disk_cache = self._load_disk_cache()
        self.stats = {"hits": 0, "misses": 0, "disk_hits": 0}
    
    def _load_disk_cache(self) -> Dict:
        """Load disk cache with size limit"""
        try:
            if os.path.exists(self.disk_cache_path):
                with open(self.disk_cache_path, 'r', encoding='utf-8') as f:
                    cache = json.load(f)
                    # Limit disk cache size
                    if len(cache) > 50000:
                        # Keep most recent half
                        items = list(cache.items())
                        cache = dict(items[-25000:])
                    return cache
        except Exception:
            pass
        return {}
    
    def _save_disk_cache(self):
        """Persist to disk periodically"""
        try:
            with open(self.disk_cache_path, 'w', encoding='utf-8') as f:
                json.dump(self.disk_cache, f, ensure_ascii=False, indent=0)
        except Exception as e:
            print(f"Cache save error: {e}")
    
    def get(self, key: str) -> Optional[str]:
        # Check memory first
        if HAS_CACHETOOLS:
            result = self.memory_cache.get(key)
            if result:
                self.stats["hits"] += 1
                return result
        else:
            if key in self.memory_cache:
                self.stats["hits"] += 1
                return self.memory_cache[key]
        
        # Check disk
        if key in self.disk_cache:
            result = self.disk_cache[key]
            # Promote to memory
            self._set_memory(key, result)
            self.stats["disk_hits"] += 1
            return result
        
        self.stats["misses"] += 1
        return None
    
    def _set_memory(self, key: str, value: str):
        """Set value in memory cache"""
        if HAS_CACHETOOLS:
            self.memory_cache[key] = value
        else:
            self.memory_cache[key] = value
            self._memory_keys.append(key)
            # Evict oldest if needed
            while len(self._memory_keys) > self._max_memory:
                oldest = self._memory_keys.pop(0)
                self.memory_cache.pop(oldest, None)
    
    def set(self, key: str, value: str):
        """Set value in both caches"""
        self._set_memory(key, value)
        self.disk_cache[key] = value
        
        # Save disk cache occasionally (every 100 new entries)
        if len(self.disk_cache) % 100 == 0:
            self._save_disk_cache()
    
    def get_stats(self) -> Dict:
        total = self.stats["hits"] + self.stats["misses"]
        hit_rate = (self.stats["hits"] / total * 100) if total > 0 else 0
        return {
            **self.stats,
            "total_requests": total,
            "hit_rate": f"{hit_rate:.1f}%",
            "memory_size": len(self.memory_cache),
            "disk_size": len(self.disk_cache)
        }

# Global cache instance
translation_cache = TranslationCache()

# ============================================================================
# TEXT PROCESSING UTILITIES
# ============================================================================

def get_cache_key(text: str, domain: str) -> str:
    """Generate compact cache key"""
    return hashlib.sha256(f"{domain}:{text}".encode()).hexdigest()[:24]

def contains_chinese(text: str) -> bool:
    """Fast check for Chinese characters"""
    # Early length check
    if len(text) < 2:
        return False
    # Use any() with generator for early exit
    return any('\u4e00' <= char <= '\u9fff' for char in text)

def should_translate(text: str) -> bool:
    """Optimized version with early Chinese detection"""
    if not text:
        return False
    
    text = text.strip()
    if len(text) < 2:
        return False
    
    # Fast path: must contain Chinese
    if not contains_chinese(text):
        return False
    
    # Skip common non-translatable patterns
    if text.isdigit():
        return False
    if text.startswith(('http://', 'https://')):
        return False
    if '@' in text and '.' in text.split('@')[-1]:
        return False
    
    return True

# ============================================================================
# BATCH TRANSLATION SYSTEM
# ============================================================================

class BatchTranslator:
    """Handles batching and deduplication of translations"""
    
    def __init__(self, client, batch_size: int = BATCH_SIZE):
        self.client = client
        self.batch_size = batch_size
        self.stats = {"api_calls": 0, "texts_translated": 0, "cached": 0}
    
    def translate_batch(self, texts: List[str], domain: str = "general") -> List[str]:
        """
        Translate a batch of texts with deduplication and caching.
        Returns translations in original order (including duplicates).
        """
        if not texts:
            return []
        
        # Build mapping for deduplication
        unique_texts = []
        text_to_indices = {}
        cache_results = {}
        
        for i, text in enumerate(texts):
            if not should_translate(text):
                cache_results[i] = text
                continue
            
            cache_key = get_cache_key(text, domain)
            cached = translation_cache.get(cache_key)
            
            if cached is not None:
                cache_results[i] = cached
                self.stats["cached"] += 1
                continue
            
            if text not in text_to_indices:
                text_to_indices[text] = []
                unique_texts.append(text)
            text_to_indices[text].append(i)
        
        # Translate in batches
        all_translations = {}
        for i in range(0, len(unique_texts), self.batch_size):
            batch = unique_texts[i:i + self.batch_size]
            translations = self._call_api_batch(batch, domain)
            
            for text, translation in zip(batch, translations):
                all_translations[text] = translation
                # Cache the result
                translation_cache.set(get_cache_key(text, domain), translation)
        
        # Build final result list
        results = []
        for i, text in enumerate(texts):
            if i in cache_results:
                results.append(cache_results[i])
            else:
                results.append(all_translations[text])
        
        self.stats["texts_translated"] += len(unique_texts)
        return results
    
    def _call_api_batch(self, texts: List[str], domain: str) -> List[str]:
        """Make a single API call for multiple texts"""
        if not texts:
            return []
        
        domain_config = DOMAINS.get(domain, DOMAINS["general"])
        prompt = domain_config["compact"] if USE_COMPACT_PROMPTS else domain_config["prompt"]
        
        # Format numbered texts
        numbered = "\n".join([f"[{i}] {t}" for i, t in enumerate(texts)])
        
        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": f"{prompt} Translate each numbered text. Respond ONLY with translations in same format."
                    },
                    {
                        "role": "user",
                        "content": f"Translate:\n{numbered}"
                    }
                ],
                temperature=0.3,
                max_tokens=len(texts) * 200
            )
            
            self.stats["api_calls"] += 1
            return self._parse_numbered_response(response.choices[0].message.content, len(texts))
            
        except Exception as e:
            print(f"Batch API error: {e}")
            # Fallback: return originals
            return texts
    
    def _parse_numbered_response(self, response: str, expected_count: int) -> List[str]:
        """Parse numbered translations from API response"""
        translations = []
        
        # Try to extract [N] format
        pattern = r'\[(\d+)\]\s*(.*?)(?=\[\d+\]|$)'
        matches = re.findall(pattern, response, re.DOTALL)
        
        if matches:
            # Sort by index and extract translations
            indexed = {int(idx): text.strip() for idx, text in matches}
            for i in range(expected_count):
                translations.append(indexed.get(i, ""))
        else:
            # Fallback: split by newlines
            lines = [l.strip() for l in response.split('\n') if l.strip()]
            translations = lines[:expected_count]
            # Pad if needed
            while len(translations) < expected_count:
                translations.append("")
        
        return translations

# ============================================================================
# DOCUMENT PROCESSING
# ============================================================================

def extract_translatable_elements(doc: Document):
    """
    Generator that yields all translatable elements from a document.
    Memory-efficient: doesn't build full list.
    """
    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            yield ('paragraph', para)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        yield ('cell', para)
    
    # Headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        yield ('header', para)
        
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    if para.text.strip():
                        yield ('footer', para)

def translate_docx_optimized(file_buffer: bytes, domain: str = "general") -> Tuple[BytesIO, Dict]:
    """
    Memory-optimized DOCX translation with batching.
    """
    start_time = time.time()
    stats = {
        "paragraphs": 0,
        "tables": 0,
        "cells": 0,
        "headers": 0,
        "footers": 0,
        "api_calls": 0,
        "cached": 0,
        "skipped": 0
    }
    
    # Load document
    doc = Document(BytesIO(file_buffer))
    
    # Collect all translatable elements
    elements = list(extract_translatable_elements(doc))
    
    # Filter to only those needing translation
    translatable = []
    for elem_type, elem in elements:
        text = elem.text.strip()
        if should_translate(text):
            translatable.append((elem_type, elem, text))
        else:
            stats["skipped"] += 1
        
        # Update stats
        if elem_type == 'paragraph':
            stats["paragraphs"] += 1
        elif elem_type == 'cell':
            stats["cells"] += 1
        elif elem_type == 'header':
            stats["headers"] += 1
        elif elem_type == 'footer':
            stats["footers"] += 1
    
    stats["tables"] = len(doc.tables)
    
    # Batch translate
    translator = BatchTranslator(get_deepseek_client(), batch_size=BATCH_SIZE)
    texts_to_translate = [t[2] for t in translatable]
    translations = translator.translate_batch(texts_to_translate, domain)
    
    # Apply translations
    for (elem_type, elem, original), translation in zip(translatable, translations):
        if elem.runs:
            elem.runs[0].text = translation
            for run in elem.runs[1:]:
                run.text = ""
    
    # Update stats
    api_stats = translator.stats
    stats["api_calls"] = api_stats["api_calls"]
    stats["cached"] = api_stats["cached"]
    stats["processing_time"] = round(time.time() - start_time, 2)
    
    # Save to buffer
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output, stats

def translate_xlsx_optimized(file_buffer: bytes, domain: str = "general") -> Tuple[BytesIO, Dict]:
    """
    Memory-optimized XLSX translation using read-only mode.
    """
    from openpyxl import load_workbook
    
    start_time = time.time()
    stats = {"sheets": 0, "cells": 0, "api_calls": 0, "cached": 0}
    
    # Use read-only for input if possible, but we need to modify
    # So we'll process in batches
    wb = load_workbook(BytesIO(file_buffer))
    
    all_cells = []
    cell_refs = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        stats["sheets"] += 1
        
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if should_translate(cell.value):
                        all_cells.append(cell.value)
                        cell_refs.append(cell)
                        stats["cells"] += 1
    
    # Batch translate
    translator = BatchTranslator(get_deepseek_client(), batch_size=BATCH_SIZE)
    translations = translator.translate_batch(all_cells, domain)
    
    # Apply translations
    for cell, translation in zip(cell_refs, translations):
        cell.value = translation
    
    api_stats = translator.stats
    stats["api_calls"] = api_stats["api_calls"]
    stats["cached"] = api_stats["cached"]
    stats["processing_time"] = round(time.time() - start_time, 2)
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    wb.close()
    
    return output, stats

# ============================================================================
# FLASK ROUTES
# ============================================================================

@app.route('/health', methods=['GET'])
def health():
    cache_stats = translation_cache.get_stats()
    return jsonify({
        "status": "ok",
        "service": "格不丢翻译 API (Optimized)",
        "version": "2.0.0",
        "domains": list(DOMAINS.keys()),
        "config": {
            "batch_size": BATCH_SIZE,
            "max_concurrent": MAX_CONCURRENT,
            "use_compact_prompts": USE_COMPACT_PROMPTS,
            "has_cachetools": HAS_CACHETOOLS
        },
        "cache": cache_stats
    })

@app.route('/domains', methods=['GET'])
def get_domains():
    return jsonify({
        "domains": {k: v["name"] for k, v in DOMAINS.items()}
    })

@app.route('/translate', methods=['POST'])
def translate():
    start_time = time.time()
    
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    domain = request.form.get('domain', 'general')
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    is_docx = file.filename.endswith('.docx')
    is_xlsx = file.filename.endswith('.xlsx')
    
    if not is_docx and not is_xlsx:
        return jsonify({"error": "Only .docx and .xlsx files are supported"}), 400
    
    # Read file
    file_buffer = file.read()
    file_size = len(file_buffer)
    
    # Check size (16MB limit)
    if file_size > 16 * 1024 * 1024:
        return jsonify({"error": "File too large. Maximum size is 16MB"}), 413
    
    try:
        if is_docx:
            output, stats = translate_docx_optimized(file_buffer, domain)
            output_filename = file.filename.replace('.docx', '_EN.docx')
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            output, stats = translate_xlsx_optimized(file_buffer, domain)
            output_filename = file.filename.replace('.xlsx', '_EN.xlsx')
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        total_time = round(time.time() - start_time, 2)
        
        # Add headers with stats
        response = send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=output_filename
        )
        response.headers['X-Processing-Time'] = str(total_time)
        response.headers['X-API-Calls'] = str(stats.get('api_calls', 0))
        response.headers['X-Cached'] = str(stats.get('cached', 0))
        
        return response
    
    except Exception as e:
        print(f"Translation error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/stats', methods=['GET'])
def get_stats():
    return jsonify(translation_cache.get_stats())

@app.route('/cache/clear', methods=['POST'])
def clear_cache():
    translation_cache.memory_cache.clear()
    translation_cache.disk_cache.clear()
    translation_cache._save_disk_cache()
    return jsonify({"status": "cache cleared"})

# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    print(f"Starting Gebudiu API (Optimized)")
    print(f"Batch size: {BATCH_SIZE}")
    print(f"Max concurrent: {MAX_CONCURRENT}")
    print(f"Compact prompts: {USE_COMPACT_PROMPTS}")
    print(f"CacheTools available: {HAS_CACHETOOLS}")
    
    app.run(debug=True, host='0.0.0.0', port=5000)
